# -*- coding: utf-8 -*-
"""
Created on Fri Jun 22 20:16:05 2018

@author: Administrator
"""

import re
import xlrd
from docx import Document
import os
import xlwt
import time,random
import requests


def ifhave(ttt):
  num=ttt
  if os.path.exists("题库.xls"):
    if os.path.exists("问题(有题库).docx"):
        path = "问题(有题库).docx"
        document = Document(path)
        b=''
        work=xlwt.Workbook()
        sheet =work.add_sheet("答案")
        style = xlwt.easyxf('font: bold 1, color red;')
        sheet.row(0).height = 256 *3
        sheet.col(3).width = 256 *50
        sheet.col(0).width = 256 *10
        sheet.col(1).width = 256 *10
        sheet.col(2).width = 256 *10
        sheet.write(0,3,'题目',style)
        sheet.write(0,1,'答案',style)
        sheet.write(0,0,'题目号',style)
        sheet.write(0,2,'备注',style)
                 
        for paragraph in document.paragraphs:
            b=b+paragraph.text.replace('\n','')
            
        
        data=xlrd.open_workbook('题库.xls')
        
        if num=='1':
            table1=data.sheets()[0]
            col1= table1.col_values(0)
            col17= table1.col_values(7)
            box1=[]
            box17=[]
            for i1 in col1:
                box1.append(i1.replace(' ',''))
            for i17 in col17:
                box17.append(i17)
            box_question=[]
            box_answer=[]
            for item in box1:
                box_question.append(item)
            for item in box17:
                box_answer.append(item)
        elif num=='2':
            table1=data.sheets()[0]
            col1= table1.col_values(0)
            col17= table1.col_values(7)
            
            table2=data.sheets()[1]
            col2= table2.col_values(0)
            col27= table2.col_values(2)
            box1=[]
            box17=[]
            for i1 in col1:
                box1.append(i1.replace(' ',''))
            for i17 in col17:
                box17.append(i17)
            box2=[]
            box27=[]
            for i2 in col2:
                box2.append(i2.replace(' ',''))
            for i27 in col27:
                box27.append(i27)
            box_question=[]
            box_answer=[]
            for item in box1:
                box_question.append(item)
            for item in box2:
                box_question.append(item)
            for item in box17:
                box_answer.append(item)
            for item in box27:
                box_answer.append(item)
        elif num=='3':
            table1=data.sheets()[0]
            col1= table1.col_values(0)
            col17= table1.col_values(7)
            
            table2=data.sheets()[1]
            col2= table2.col_values(0)
            col27= table2.col_values(7)
            
            table3=data.sheets()[2]
            col3=table3.col_values(0)
            col32= table3.col_values(2)
            box1=[]
            box17=[]
            for i1 in col1:
                box1.append(i1.replace(' ',''))
            for i17 in col17:
                box17.append(i17)
            box2=[]
            box27=[]
            for i2 in col2:
                box2.append(i2.replace(' ',''))
            for i27 in col27:
                box27.append(i27)
            box3=[]
            box32=[]
            for i3 in col3:
                box3.append(i3.replace(' ',''))
            for i32 in col32:
                box32.append(i32)
            box_question=[]
            box_answer=[]
            for item in box1:
                box_question.append(item)
            for item in box2:
                box_question.append(item)
            for item in box3:
                box_question.append(item)
            for item in box17:
                box_answer.append(item)
            for item in box27:
                box_answer.append(item)
            for item in box32:
                box_answer.append(item)
        elif num=='4':
            table1=data.sheets()[0]
            col1= table1.col_values(0)
            col17= table1.col_values(7)
            
            table2=data.sheets()[1]
            col2= table2.col_values(0)
            col27= table2.col_values(7)
            
            box1=[]
            box17=[]
            for i1 in col1:
                box1.append(i1.replace(' ',''))
            for i17 in col17:
                box17.append(i17)
            box2=[]
            box27=[]
            for i2 in col2:
                box2.append(i2.replace(' ',''))
            for i27 in col27:
                box27.append(i27)
            for item in box1:
                box_question.append(item)
            for item in box2:
                box_question.append(item)
            for item in box17:
                box_answer.append(item)
            for item in box27:
                box_answer.append(item)
        #print(box_question)
        n=0
        a=re.compile('(\d+)\.(.*?)A').findall(b)
        #print(a)
        for i in range(len(a)):
            n=n+1
            m=0
            k=-1
            yuanques=a[i][1].replace(' ','').replace('\n','').replace('.','').replace('：','').replace('（','').replace('？','').replace('！','').replace('）','').replace('。','')
            #print(yuanques)
            #print(a[i][1].replace(' ',''))
            for j in box_question:
                k=k+1
                if  yuanques in j.replace(' ','').replace('.','').replace('：','').replace('\n','').replace('（','').replace('？','').replace('！','').replace('）','').replace('。',''):
                    m=1
                    print('第'+str(n)+'题: '+str(j)+'\n'+'选:'+box_answer[k])
                    sheet.write(n,3,j)
                    sheet.write(n,1,box_answer[k])
                    break
            if m==0:
                print('题库中无此题目，为了方便他人，请将这道题加入题库。谢谢')
                sheet.write(n,3,yuanques)
                sheet.write(n,1,'无答案')
                sheet.write(n,2,'自行百度')
            sheet.write(n,0,'第'+str(n)+'题')
        work.save('生成的答案(有题库版).xls')
    else:
        print('缺失文件！！！请在本程序目录建立一个“问题(有题库).docx”的文件')
        print('输入任意字符回车后退出\n')
        input()
  else:
    print('缺失文件！！！请把题库放在和该程序相同的目录下，并命名为“题库.xls”')
    print('输入任意字符回车后退出\n')
    input()
def nothave():
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 6.1; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/53.0.2785.143 Safari/537.36'}
    path = "问题(没有题库).docx"
    document = Document(path)
    b=''
    for paragraph in document.paragraphs:
            b=b+paragraph.text.replace('\n','')
    a=re.compile('(\d+)\.(.*?)A').findall(b)
    surl='https://www.asklib.com/s/'
    work=xlwt.Workbook()
    sheet =work.add_sheet("答案")
    style = xlwt.easyxf('font: bold 1, color red;')
    sheet.row(0).height = 256 *3
    sheet.col(3).width = 256 *80
    sheet.col(4).width = 256 *80
    sheet.col(0).width = 256 *10
    sheet.col(1).width = 256 *10
    sheet.col(2).width = 256 *20
    sheet.row(1).height = 256 *3
    sheet.write(1,3,'测试的题目',style)
    sheet.write(1,1,'参考答案',style)
    sheet.write(1,0,'题目号',style)
    sheet.write(1,4,'网络上找到的相似的题目',style)
    sheet.write(1,2,'参考答案对应选项',style)
    sheet.write(0,0,'由于没有题库，下列答案是在互联网上找到的，仅供参考，不一定正确！！！！！如有疑问，请自行百度',style)
    n=2
    for i in range(len(a)):
        print('正在下载第'+str(n-1)+'题')
        ques=str(a[i][1].replace(' ','').replace('\n',''))
        url=surl+ques
        #print(url)
        r=requests.get(url,headers=headers)
        if str(r.status_code)=='200':
          yq=re.compile('<h2 class="F18 ti2m LH_40">(.*?)</h2>').findall(r.text)[0].replace('<em>','').replace('</em>','').replace('_','')
          url1='https://www.asklib.com'
          u=re.compile('<a href="(.*?)">参考解析</a>').findall(r.text)[0]
          url1=str(url1)+str(u)
          rr=requests.get(url1,headers=headers)
          test=re.compile('<h1 class="F18 ti2m LH_40">(.*?)</h1>').findall(rr.text)[0]
          if '[单选]' in test:
            an=re.compile('<div class="listtip">参考答案：(.*?)</div>').findall(rr.text)[0]
            ans=re.compile('<p class="F18 ml2m LH_40">(.*?)<br /></p>').findall(rr.text)[0].replace('<br />','').replace('<br/>','').replace('<p>','')
            ans=ans+'哈'
            if 'D' in ans and len(an)==1:
                #print(ans)
                A=re.compile('A(.*?)B').findall(ans)[0].replace('.','').replace(' ','')
                B=re.compile('B(.*?)C').findall(ans)[0].replace('.','').replace(' ','')
                C=re.compile('C(.*?)D').findall(ans)[0].replace('.','').replace(' ','')
                D=re.compile('D(.*?)哈').findall(ans)[0].replace('.','').replace(' ','')
                if str(an)=='A':
                    answer=str(A)
                elif str(an)=='B':
                    answer=str(B)
                elif str(an)=='C':
                    answer=str(C)
                elif str(an)=='D':
                    answer=str(D)
                else:
                    pass
                print('选'+an)
                print(answer)
                sheet.write(n,0,'第'+str(n-1)+'题')
                sheet.write(n,1,an)
                sheet.write(n,2,answer)
                sheet.write(n,3,ques)
                sheet.write(n,4,yq)
                time.sleep(random.choice([0.2,0.5,1]))
            else:
                sheet.write(n,0,'第'+str(n-1)+'题')
                sheet.write(n,1,'没找到答案')
                sheet.write(n,2,'无')
                sheet.write(n,3,ques)
                sheet.write(n,4,'')
                #time.sleep(random.choice([0.2,0.5,1]))
            
        else:
                sheet.write(n,0,'第'+str(n-1)+'题')
                sheet.write(n,1,'没找到答案')
                sheet.write(n,2,'无')
                sheet.write(n,3,ques)
                sheet.write(n,4,'')
                time.sleep(random.choice([0.2,0.5,1]))
        n=n+1
    work.save('生成的答案(无题库版).xls')
judge=str(input('是否有题库？（输入“有”或者“没有”后按下回车键）【备注：无题库仅适合刷单选题，多选和判断是找不到答案的】\n'))
if judge=='有':
    test=str(input('请题目类型数（仅单选输入1，单选加判断输入2，单选加多选加判断输入3，单选加多选输入4）,然后按下回车键确认！！\n'))
    ifhave(test)
elif judge=='没有':
    if os.path.exists("问题(没有题库).docx"):
        nothave()
    else:
        print('缺失文件！！！请在本程序目录建立一个“问题(没有题库).docx”的文件\n')
        print('输入任意字符回车后退出\n')
        input() 
else:
    print('您的输入有误，请退出程序重新打开输入正确的判断指令（输入任意字符回车后退出）\n')
    input()      